# -*- coding: utf-8 -*-
from copy import deepcopy
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
SRC_PATH = ROOT / "tmp" / "win.docx"
DEFAULT_OUT_PATH = ROOT / "tmp" / "technical_specification_rtrace.docx"


def add_paragraph(doc: Document, text: str = "", style: str = "Normal", align=None):
    paragraph = doc.add_paragraph(style=style)
    if text:
        paragraph.add_run(text)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_heading_1(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.add_run(text)
    return paragraph


def add_heading_2(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.add_run(text)
    return paragraph


def add_body(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(text)
    return paragraph


def add_list(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Paragraph")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(text)
    return paragraph


def build_doc(out_path: Path):
    template = Document(str(SRC_PATH))
    source = Document(str(SRC_PATH))

    body = template._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    add_paragraph(
        template,
        "Государственное бюджетное общеобразовательное учреждение города Москвы «Школа № 1533 «ЛИТ»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for _ in range(5):
        add_paragraph(template, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(template, "Техническое задание", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        template,
        "к проектной работе направления «Blue Team» по теме: "
        "«Программа для анализа ELF бинарных файлов на признаки вредоносного программного обеспечения»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for _ in range(5):
        add_paragraph(template, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    if source.tables:
        body.insert(len(body) - 1, deepcopy(source.tables[0]._element))

    for _ in range(4):
        add_paragraph(template, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    page_break = add_paragraph(template, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break.add_run().add_break(WD_BREAK.PAGE)

    add_heading_1(template, "1. ОБЩИЕ ПОЛОЖЕНИЯ")
    add_heading_2(template, "1.1 Наименование проекта")
    add_body(
        template,
        "Наименование проектного продукта: rtrace — программа для анализа "
        "ELF-бинарных файлов Linux x86-64 на признаки вредоносной активности "
        "в изолированной среде.",
    )
    add_heading_2(template, "1.2 Основание для разработки")
    add_body(
        template,
        "Настоящее техническое задание разработано в рамках проектной работы "
        "направления Blue Team. Основанием для разработки является "
        "необходимость создать воспроизводимый инструмент первичного анализа "
        "подозрительных ELF-файлов, который можно безопасно использовать в "
        "учебном и исследовательском стенде.",
    )
    add_heading_2(template, "1.3 Назначение разработки")
    add_body(
        template,
        "Разрабатываемый продукт предназначен для первичного динамического "
        "анализа ELF-файлов под Linux в изолированной виртуальной машине. "
        "Утилита должна автоматически выявлять признаки вредоносной "
        "активности по YARA-сигнатурам в исполняемом файле и памяти процесса, "
        "сохранять доказательные артефакты и подготавливать результат для "
        "последующего разбора.",
    )
    add_heading_2(template, "1.4 Цель разработки")
    add_body(
        template,
        "Цель разработки — создать инструмент rtrace, который запускает "
        "анализ в изолированной Linux-среде, автоматически находит IoC и TTP "
        "по YARA-правилам, формирует структурированные JSON-артефакты, "
        "сопоставляет находки с MITRE ATT&CK и при необходимости "
        "останавливает процесс после первого нового срабатывания.",
    )

    add_heading_1(template, "2. ТРЕБОВАНИЯ К ПРОЕКТНОМУ ПРОДУКТУ")
    add_heading_2(template, "2.1 Функциональные требования")
    add_body(
        template,
        "Продукт должен поддерживать запуск гостевой утилиты rtrace с "
        "параметрами `--rules-dir`, `--samples-dir`, `--pid`, "
        "`--artifacts-dir`, `--scan-interval-ms`, `--max-region-bytes`, "
        "`--max-total-bytes`, `--once`, `--verbose`, `--stop-on-hit`, "
        "`--save-maps`, `--dump-regions`.",
    )
    add_list(
        template,
        "Программа должна поддерживать три режима выбора целей анализа: "
        "процессы из каталога `--samples-dir`, процессы по указанному `--pid` "
        "и все несистемные процессы, если селектор не задан.",
    )
    add_list(
        template,
        "Программа должна выполнять сканирование по каналам `FILE` "
        "(исполняемый файл процесса), `MEM` (читаемые регионы памяти) и "
        "`REG` (память по указателям из CPU-регистров).",
    )
    add_list(
        template,
        "При каждом новом совпадении программа должна автоматически создавать "
        "JSON-артефакт в каталоге `artifacts`.",
    )
    add_list(
        template,
        "В артефакте должны сохраняться контекст процесса и доказательная "
        "информация: `timestamp_ms`, `pid`, `ppid`, `uid`, `exe`, `cmdline`, "
        "`rule`, `class`, `channel`, `cpu_register`, `address`, "
        "`match_string`, `match_offset`, `matched_bytes_hex`, "
        "`matched_bytes_ascii`, а также сопоставление с MITRE ATT&CK.",
    )
    add_list(
        template,
        "Программа должна поддерживать режим активной защиты: при включенном "
        "флаге `--stop-on-hit` процесс, для которого получены новые "
        "срабатывания, должен быть автоматически остановлен.",
    )
    add_list(
        template,
        "Результаты анализа должны быть пригодны для последующего просмотра в "
        "WebUI по временной шкале и по связям между процессами.",
    )

    add_heading_2(template, "2.2 Нефункциональные требования")
    add_list(template, "Реализация должна быть выполнена на языке Rust с модульной архитектурой.")
    add_list(template, "Запуск стенда должен быть воспроизводимым в среде Windows + WSL + QEMU.")
    add_list(
        template,
        "Анализируемые ELF-файлы не должны запускаться на хостовой системе "
        "при использовании режима виртуализации.",
    )
    add_list(
        template,
        "Программа должна сохранять работоспособность при ошибках чтения "
        "`/proc`, завершении процесса во время анализа и отсутствии части "
        "ожидаемых данных.",
    )
    add_list(
        template,
        "Набор YARA-правил должен расширяться без изменения основной логики программы.",
    )
    add_list(
        template,
        "Программа должна предоставлять CLI-интерфейс и режим `--help` с примерами запуска.",
    )

    add_heading_2(template, "2.3 Требования к безопасности и изоляции")
    add_body(
        template,
        "Запуск подозрительных ELF-файлов должен выполняться в изолированной "
        "виртуальной машине Linux под QEMU. Для взаимодействия между хостом и "
        "гостевой системой должны использоваться только выделенные каталоги "
        "`/rules`, `/samples` и `/artifacts`. Каталоги с правилами и "
        "образцами монтируются в режиме только для чтения, а каталог "
        "артефактов — в режиме записи. Решение не должно требовать запуска "
        "анализируемых бинарных файлов непосредственно на хостовой системе.",
    )

    add_heading_2(template, "2.4 Требования к результатам анализа")
    add_body(
        template,
        "Результатом работы продукта должен быть воспроизводимый набор "
        "артефактов, по которым можно подтвердить факт совпадения сигнатуры, "
        "определить канал обнаружения, связать событие с конкретным процессом "
        "и интерпретировать его в терминах TTP и MITRE ATT&CK. Отдельно "
        "должны поддерживаться демонстрационные сценарии, в которых часть "
        "сигнатур проявляется только в памяти процесса в ходе выполнения.",
    )

    add_heading_1(template, "3. КРИТЕРИИ ПРИЕМКИ")
    add_body(
        template,
        "Техническое задание считается выполненным, если выполнены следующие "
        "критерии приемки.",
    )
    add_list(template, "Утилита успешно собирается и запускается в целевом стенде Windows + WSL + QEMU.")
    add_list(
        template,
        "Программа корректно загружает активный набор YARA-правил и переходит к анализу без изменения исходного кода.",
    )
    add_list(
        template,
        "На контрольном наборе демонстрационных ELF-образцов фиксируются "
        "срабатывания как минимум по каналам `FILE` и `MEM`.",
    )
    add_list(
        template,
        "В артефактах `meta.json` присутствуют обязательные поля процесса, "
        "канала, адреса, смещения и совпавших байтов.",
    )
    add_list(
        template,
        "Режим `--stop-on-hit` подтверждается практическим прогоном, в "
        "котором анализируемый процесс автоматически останавливается после "
        "нового срабатывания.",
    )
    add_list(
        template,
        "На демонстрационном наборе сэмплов фиксируется не менее 15 "
        "уникальных сигнатур YARA.",
    )
    add_list(
        template,
        "Полученные артефакты корректно открываются в WebUI и отображаются по "
        "временной шкале и по связям процессов.",
    )

    add_heading_1(template, "4. ПОРЯДОК ТЕСТИРОВАНИЯ")
    add_heading_2(template, "4.1 Методика тестирования")
    add_body(
        template,
        "Тестирование должно проводиться в изолированной виртуальной машине "
        "Ubuntu под QEMU с примонтированными каталогами `/rules`, `/samples` "
        "и `/artifacts`. Проверка включает модульные тесты, контрольные "
        "end-to-end прогоны и просмотр полученных артефактов в WebUI.",
    )
    add_heading_2(template, "4.2 Контрольные сценарии")
    add_list(
        template,
        "Сценарий статического и динамического анализа демонстрационного "
        "ELF-образца с сигнатурами Linux-ориентированных TTP.",
    )
    add_list(
        template,
        "Сценарий runtime-детекта, в котором часть строк раскрывается только "
        "в памяти дочернего процесса.",
    )
    add_list(
        template,
        "Сценарий активной защиты с использованием `--stop-on-hit`.",
    )
    add_heading_2(template, "4.3 Метрики тестирования")
    add_list(template, "Количество сформированных снапшотов.")
    add_list(template, "Общее число срабатываний и число уникальных правил.")
    add_list(template, "Распределение срабатываний по каналам `FILE`, `MEM`, `REG`.")
    add_list(template, "Время до появления первого артефакта.")
    add_list(
        template,
        "Полнота полей в `meta.json` и корректность сопоставления с MITRE ATT&CK.",
    )

    add_heading_1(template, "5. СОСТАВ РЕЗУЛЬТАТОВ РАЗРАБОТКИ")
    add_body(
        template,
        "По итогам выполнения технического задания должны быть подготовлены "
        "исходный код проекта, набор активных YARA-правил, скрипты "
        "подготовки и запуска стенда, демонстрационные ELF-образцы, WebUI "
        "для просмотра артефактов, а также пояснительная документация с "
        "результатами тестирования и описанием архитектуры решения.",
    )

    template.save(str(out_path))


if __name__ == "__main__":
    out_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUT_PATH
    build_doc(out_path)
